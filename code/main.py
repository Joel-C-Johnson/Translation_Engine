import os
import uuid
import json
import psycopg2
from functools import wraps
from datetime import datetime, timedelta
import datetime
import scrypt
import requests
import jwt
from flask import Flask, request, session, redirect
from flask import g
from flask_cors import CORS, cross_origin
import nltk
import re
import xlrd
from xlrd import open_workbook
import json
import ast
import flask
import pyexcel
import logging
from docx import Document
import pdb

# @app.route("/v1/save", methods=["POST", "GET"])      #------------------To save file in db------------------------#
# def save():
#     # content = request.files['content']
#     # fname = content.read()
#     f = open('ff.docx','rb')
#     dat = f.read()
#     binary = psycopg2.Binary(dat)
#     connection = get_db()
#     cursor = connection.cursor()
#     cursor.execute("INSERT INTO save (file) VALUES (%s)", (binary,))
#     cursor.close()
#     connection.commit()
#     return 'success'

@app.route("/v1/file", methods=["POST", "GET"])           #------------------To fetch file from db------------------------#
def file():
    source_id = 1
    connection = get_db()
    cursor = connection.cursor()
    cursor.execute("SELECT  file FROM other_sources WHERE id = %s ", (source_id,))
    file1 = cursor.fetchone()
    open("file.docx", 'wb').write(file1[0])
    cursor.close()
    return 'done'

@app.route("/v1/source", methods=["POST", "GET"])
# @check_token
def source_word():
    content = request.files['source_file']
    fname = content.read()
    name = request.form["name"]
    file_type = request.form["file_type"]
    language = request.form["language"]
    connection = get_db()
    cursor = connection.cursor()
    cursor.execute("SELECT id FROM other_sources WHERE language = %s AND name = %s AND file_type = %s", (language, name, file_type))
    s_id = cursor.fetchone()
    if s_id:
        return '{"success":false, "message":"File already exists.Upload new file."}'
    else:
        with open("source.docx", "wb") as o:   # TODO: find a way to directly read from fname
            o.write(fname)
        f = open('source.docx','rb')
        dat = f.read()
        binary = psycopg2.Binary(dat)
        cursor.execute("INSERT INTO other_sources (name, language, file, file_type) VALUES (%s, %s, %s, %s)", (name, language, binary, file_type))
        cursor.close()
        connection.commit()
        doc = Document('source.docx')
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        fullText = "\n".join(fullText)
        remove_punct = re.sub(r'([!"#$%&\\\'\(\)\*\+,\.\/:;<=>\?\@\[\]^_`{|\}~\”\“\‘\’।0123456789])', '', fullText)
        token_list = set(nltk.word_tokenize(remove_punct))
        connection = get_db()
        cursor = connection.cursor()
        cursor.execute("SELECT id FROM other_sources WHERE language = %s AND name = %s AND file_type = %s", (language, name, file_type))
        source_id = cursor.fetchone()
        for t in list(token_list):
            cursor.execute("INSERT INTO other_tokens (name, tokens, source_id) VALUES (%s, %s, %s)", (name, t, source_id[0]))
        cursor.close()
        connection.commit()
        filename = "source.docx"
        if os.path.exists(filename):
            os.remove(filename)
        return '{"success":false, "message":"File has been uploaded successfully"}'

@app.route("/v1/downloadtokens", methods=["POST", "GET"])               #------------------download tokens in excel file------------------------#
# @check_token
def downloadtokens():
    language = request.form["language"]
    name = request.form["name"]
    file_type = request.form["file_type"]
    targetlang = request.form["targetlang"]
    connection = get_db()
    cursor = connection.cursor()
    cursor.execute("SELECT id FROM other_sources WHERE language = %s AND name = %s AND file_type = %s ", (language, name, file_type))
    source_id = cursor.fetchone()
    if not source_id:
        return '{"success":false, "message":"Source is not available. Upload source."}'
    else:
            translatedtokenlist = []
            cursor.execute("SELECT  token FROM other_tokentranslations WHERE translated_token IS NOT NULL AND  targetlang = %s AND source_id = %s", (targetlang, source_id[0]))
            translatedtoken = cursor.fetchall()
            for tk in translatedtoken:
                translatedtokenlist.append(tk[0])
            toknwords =[]
            cursor.execute("SELECT tokens FROM other_tokens WHERE source_id =%s AND name = %s", (source_id[0], name))
            tokens = cursor.fetchall()
            for t in tokens:
                toknwords.append(t[0])
            stoknwords = set(toknwords)
            output = stoknwords - set(translatedtokenlist)
            cursor.close()
            result = [['TOKEN', 'TRANSLATION']]
            for i in list(output):
                result.append([i])
            sheet = pyexcel.Sheet(result)
            output = flask.make_response(sheet.xlsx)
            output.headers["Content-Disposition"] = "attachment; filename =token.xlsx"
            output.headers["Content-type"] = "xlsx"
            return output

@app.route("/v1/uploadtokentranslation", methods=["POST"])                         #------------------To upload token translation in database via excel file------------------------#
# @check_token
def upload_tokens_translation():
    language = request.form["language"]
    name = request.form["name"]
    file_type = request.form["file_type"]
    targetlang = request.form["targetlang"]
    tokenwords = request.files['tokenwords']
    changes = []
    connection = get_db()
    cursor = connection.cursor()
    cursor.execute("SELECT id FROM other_sources WHERE language = %s AND name = %s AND file_type = %s", (language, name, file_type))
    source_id = cursor.fetchone()
    if not source_id:
        return '{"success":false, "message":"Unable to locate the language, version and revision number specified"}'
    exl = tokenwords.read()
    with open("tokn.xlsx", "wb") as o:
        o.write(exl)
    tokenwords = open_workbook('tokn.xlsx')
    book = tokenwords
    p = book.sheet_by_index(0)
    count = 0
    for c in range(p.nrows):                                   # to find an empty cell
        cell = p.cell(c, 1).value
        if cell:
            count = count + 1
    if count > 1:
        token_c = (token_c.value for token_c in p.col(0, 1))
        tran = (tran.value for tran in p.col(1, 1))
        data = dict(zip(token_c, tran))
        dic = ast.literal_eval(json.dumps(data))
        cursor.execute("SELECT token FROM other_tokentranslations WHERE source_id = %s AND targetlang = %s", (source_id[0], targetlang))
        transtokens = cursor.fetchall()
        if transtokens:
            token_list = []
            for i in transtokens:
                token_list.append(i[0])
            for k, v in dic.items():
                if v:
                    if k not in token_list:
                        cursor.execute("INSERT INTO other_tokentranslations (token, translated_token, targetlang, source_id) VALUES (%s, %s, %s, %s)", (k, v, targetlang, source_id[0]))
                        changes.append(v)
                    # cursor.execute("UPDATE autotokentranslations SET translated_token = %s WHERE token = %s AND source_id = %s AND targetlang = %s AND revision_num = %s", (v, k, source_id[0], targetlang, revision))
            cursor.close()
            connection.commit()
            filename = "tokn.xlsx"
            if os.path.exists(filename):
                os.remove(filename)
        else:
            for k, v in dic.items():
                if v:
                    cursor.execute("INSERT INTO other_tokentranslations (token, translated_token, targetlang, source_id) VALUES (%s, %s, %s, %s)", (k, v, targetlang, source_id[0]))
                    changes.append(v)
            cursor.close()
            connection.commit()
            filename = "tokn.xlsx"
            if os.path.exists(filename):
                os.remove(filename)
        if changes:
            # logging.warning('User \'' + str(request.email) + '\' uploaded translation of tokens successfully')
            return '{"success":true, "message":"Token translation have been uploaded successfully"}'
        else:
            # logging.warning('User \'' + str(request.email) + '\' upload of token translation unsuccessfully')
            return '{"success":false, "message":"No Changes. Existing token is already up-to-date."}'
    else:
        return '{"success":false, "message":"Tokens have no translation"}'


@app.route("/v1/translations", methods=["POST", "GET"])         #------------------download translation draft in docx file------------------------#
# @check_token
def translationdraft():
    language = request.form["language"]
    name = request.form["name"]
    file_type = request.form["file_type"]
    targetlang = request.form["targetlang"]
    changes = []
    changes1 = []
    connection = get_db()
    cursor = connection.cursor()
    tokens = {}
    cursor.execute("SELECT id FROM other_sources WHERE language = %s AND name = %s AND file_type = %s", (language, name, file_type))
    rst = cursor.fetchone()
    if not rst:
            return '{"success":false, "message":"Source is not available. Upload it"}'
    else:
        source_id = rst[0]
        cursor.execute("SELECT token, translated_token FROM other_tokentranslations WHERE targetlang = %s AND source_id = %s AND translated_token IS NOT NULL", (targetlang, source_id))
        if not cursor.fetchall():
            return "there is no translations. Select different target language"
        td = dict((t, tr) for t, tr in cursor.fetchall())
        d = sorted(td.items(), key=lambda k:  len(k[0]), reverse=True)
        doc = Document('ff.docx')
        print(d)
        for p in doc.paragraphs:
            for k, v in d:
                if k in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if k in inline[i].text:
                            text = inline[i].text.replace(k, v)
                            inline[i].text = text
        for h in doc.tables:
            for k, v in d:
                if k in h.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if k in inline[i].text:
                            text = inline[i].text.replace(k, v)
                            inline[i].text = text
        doc.save('dest1.docx')
        return 'done'

@app.route("/v1/testing", methods=["POST", "GET"])
def testing():
    doc = Document('ff.docx')
    dic = {'old' : "new", 'boy' : "girl", 'pen' : "pensil"}
    for p in doc.paragraphs:
        for k, v in dic.items():
            if k in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if k in inline[i].text:
                        text = inline[i].text.replace(k, v)
                        inline[i].text = text
    doc.save('dest1.docx')
    return '1'
